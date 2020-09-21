#!/usr/bin/python2
# -*- coding: utf-8 -*-
"""
RSA Letter-Color Training, Colored Text
O.Colizoli 2019
Text Name: Arial Black, Font Size 10
Single letter
~7.5 minutes per 25,000 words (13 letters)
Python 2.7
"""
import os, time # for paths and data
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
from IPython import embed as shell
import numpy as np
import pandas as pd
import general_parameters as gp
import docx_search_loop as sl
# -------------
# Each book needs the following format before running this script: "sub-{}_book-{}_text.docx"
# pip install python-docx
# -------------
# Character formatting is applied at the Run level. 
# Examples include font typeface and size, bold, italic, and underline.
# A Run object has a read-only font property providing access to a Font object. 
# A run's Font object provides properties for getting and setting the character formatting for that run.
# -------------
# See: https://github.com/aekoch/docx_formatting_aekoch
# -------------
# First run the search loop for the particular letters, replace colors; 
# Second, replace the entire documents font name and size, save

try:
    subject_ID = int(raw_input("Subject ID: ")) 
    book_ID = int(raw_input("Book Number: ")) 
except:
    subject_ID = int(input("Subject ID: ")) 
    book_ID = int(input("Book Number: "))
    

inFileName = os.path.join('books',"sub-{}_book-{}_text.docx".format(subject_ID,book_ID))
outFileName = os.path.join('books',"sub-{}_book-{}.docx".format(subject_ID,book_ID))

font_name = gp.font_trained
font_size = 10

## LETTER SET DEPENDENT ON SUBJECT NUMBER
trained_letters = gp.letter_sets_lower[np.mod(subject_ID,2)] # lower case, TRAINED set!!

# Get subject-specific letter-color mapping, letters always have same order, colors are permuted
DFS = pd.read_csv(os.path.join('colors','sub-{}'.format(subject_ID),'sub-{}_colors.tsv'.format(subject_ID)),sep='\t')
print(DFS)
subj_colors = [[DFS['r'][i],DFS['g'][i],DFS['b'][i]] for i in range(len(DFS))]
subj_letters = list(DFS['letter'])

letters = subj_letters
colors = subj_colors

t0 = time.time()
# First, run the letter-by-letter search and replace loop
doc = Document(inFileName)
for paragraph in doc.paragraphs:
    for i,search in enumerate(letters):
        ## print(sl.find_occurances_in_paragraph(paragraph, search))
        try:
            format_func = lambda x:x.font.color.__setattr__('rgb', RGBColor(int(colors[i][0]),int(colors[i][1]),int(colors[i][2])))
        except:
            shell()
        for start in sl.find_occurances_in_paragraph(paragraph, search):
            try:
                sl.apply_format_to_range(paragraph, start, start + len(search), format_func) 
            except:
                shell()
doc.save(outFileName)

# Second, replace the entire document's font name and size, save it
doc = Document(outFileName) # open an existing book
for paragraph in doc.paragraphs:
    for r in paragraph.runs:
        font = r.font
        font.name = font_name
        font.size = Pt(font_size)
        ##font.color.rgb = RGBColor(0, 0, 0) # set all to black 
doc.save(outFileName) # overwrite

print('It took {} minutes'.format( (time.time()-t0)/60 )) # ~7.5 min per 25,000 words
